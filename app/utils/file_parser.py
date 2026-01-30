import mimetypes
from dataclasses import dataclass, field
from pathlib import Path

import pypdf
from docx import Document as DocxDocument


@dataclass
class ParsedDocument:
    """パース結果"""

    text: str
    page_count: int | None = None
    page_texts: list[str] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)


class FileParser:
    """ファイルパーサー"""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}
    SUPPORTED_MIME_TYPES = {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
        "text/plain",
        "text/markdown",
    }

    def __init__(self):
        self._parsers = {
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
            ".doc": self._parse_docx,
            ".txt": self._parse_text,
            ".md": self._parse_text,
        }

    def parse(self, file_path: str | Path) -> ParsedDocument:
        """
        ファイルをパースしてテキストを抽出

        Args:
            file_path: ファイルパス

        Returns:
            パース結果

        Raises:
            ValueError: サポートされていないファイル形式の場合
            FileNotFoundError: ファイルが存在しない場合
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"ファイルが見つかりません: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"サポートされていないファイル形式: {ext}. "
                f"サポート形式: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )

        parser = self._parsers.get(ext)
        if parser is None:
            raise ValueError(f"パーサーが見つかりません: {ext}")

        return parser(path)

    def parse_bytes(
        self,
        content: bytes,
        filename: str,
        mime_type: str | None = None,
    ) -> ParsedDocument:
        """
        バイトデータからパース

        Args:
            content: ファイルの内容（バイト）
            filename: ファイル名
            mime_type: MIMEタイプ（オプション）

        Returns:
            パース結果
        """
        ext = Path(filename).suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            if mime_type:
                ext = self._mime_to_extension(mime_type)
            if ext not in self.SUPPORTED_EXTENSIONS:
                raise ValueError(f"サポートされていないファイル形式: {ext}")

        if ext in {".txt", ".md"}:
            text = content.decode("utf-8", errors="replace")
            return ParsedDocument(text=text, metadata={"filename": filename})

        if ext == ".pdf":
            return self._parse_pdf_bytes(content, filename)

        if ext in {".docx", ".doc"}:
            return self._parse_docx_bytes(content, filename)

        raise ValueError(f"サポートされていないファイル形式: {ext}")

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        """PDFファイルをパース"""
        page_texts: list[str] = []
        with open(path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                text = page.extract_text() or ""
                page_texts.append(text)

        full_text = "\n\n".join(page_texts)
        return ParsedDocument(
            text=full_text,
            page_count=len(page_texts),
            page_texts=page_texts,
            metadata={"filename": path.name},
        )

    def _parse_pdf_bytes(self, content: bytes, filename: str) -> ParsedDocument:
        """PDFバイトデータをパース"""
        import io

        page_texts: list[str] = []
        reader = pypdf.PdfReader(io.BytesIO(content))
        for page in reader.pages:
            text = page.extract_text() or ""
            page_texts.append(text)

        full_text = "\n\n".join(page_texts)
        return ParsedDocument(
            text=full_text,
            page_count=len(page_texts),
            page_texts=page_texts,
            metadata={"filename": filename},
        )

    def _parse_docx(self, path: Path) -> ParsedDocument:
        """Word文書をパース"""
        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)

        return ParsedDocument(
            text=full_text,
            metadata={"filename": path.name, "paragraph_count": len(paragraphs)},
        )

    def _parse_docx_bytes(self, content: bytes, filename: str) -> ParsedDocument:
        """Wordバイトデータをパース"""
        import io

        doc = DocxDocument(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)

        return ParsedDocument(
            text=full_text,
            metadata={"filename": filename, "paragraph_count": len(paragraphs)},
        )

    def _parse_text(self, path: Path) -> ParsedDocument:
        """テキストファイルをパース"""
        text = path.read_text(encoding="utf-8", errors="replace")
        return ParsedDocument(
            text=text,
            metadata={"filename": path.name},
        )

    def _mime_to_extension(self, mime_type: str) -> str:
        """MIMEタイプから拡張子を取得"""
        ext = mimetypes.guess_extension(mime_type)
        if ext:
            return ext

        mime_map = {
            "application/pdf": ".pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
            "application/msword": ".doc",
            "text/plain": ".txt",
            "text/markdown": ".md",
        }
        return mime_map.get(mime_type, "")

    @classmethod
    def is_supported(cls, filename: str, mime_type: str | None = None) -> bool:
        """ファイルがサポートされているか確認"""
        ext = Path(filename).suffix.lower()
        if ext in cls.SUPPORTED_EXTENSIONS:
            return True
        if mime_type and mime_type in cls.SUPPORTED_MIME_TYPES:
            return True
        return False

    @classmethod
    def get_mime_type(cls, filename: str) -> str | None:
        """ファイル名からMIMEタイプを推測"""
        mime_type, _ = mimetypes.guess_type(filename)
        return mime_type
